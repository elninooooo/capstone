#!/usr/bin/env python3
"""IoT 区块链数据存证系统 - 使用手册生成脚本"""
import os
import subprocess
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = BASE_DIR


def set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def generate_architecture_diagram():
    """生成系统架构图"""
    dot_src = """
    digraph architecture {
        rankdir=TB;
        fontname="Microsoft YaHei";
        node [fontname="Microsoft YaHei", shape=box, style="rounded,filled", fontsize=11];
        edge [fontname="Microsoft YaHei", fontsize=9];
        bgcolor="white";

        subgraph cluster_frontend {
            label="表现层 (Frontend)";
            style="dashed,rounded";
            color="#4A90D9";
            fontcolor="#4A90D9";
            fontsize=12;
            frontend [label="React 前端\\n(Ant Design 仪表盘)", fillcolor="#E3F2FD"];
        }

        subgraph cluster_backend {
            label="控制层 (Backend)";
            style="dashed,rounded";
            color="#66BB6A";
            fontcolor="#66BB6A";
            fontsize=12;
            flask [label="Flask REST API\\n(Python Web 服务)", fillcolor="#E8F5E9"];
            web3py [label="Web3.py\\n(区块链交互客户端)", fillcolor="#E8F5E9"];
            ipfs_client [label="IPFS 客户端\\n(文件存储服务)", fillcolor="#E8F5E9"];
        }

        subgraph cluster_blockchain {
            label="逻辑层 (Blockchain)";
            style="dashed,rounded";
            color="#FF9800";
            fontcolor="#FF9800";
            fontsize=12;
            contract [label="IoTDataStorage\\nSolidity 智能合约", fillcolor="#FFF3E0"];
            hardhat [label="Hardhat 本地链\\n(以太坊节点)", fillcolor="#FFF3E0"];
        }

        subgraph cluster_storage {
            label="数据层 (Storage)";
            style="dashed,rounded";
            color="#AB47BC";
            fontcolor="#AB47BC";
            fontsize=12;
            ipfs [label="IPFS / 本地存储\\n(内容寻址文件系统)", fillcolor="#F3E5F5"];
        }

        subgraph cluster_iot {
            label="数据源 (IoT)";
            style="dashed,rounded";
            color="#78909C";
            fontcolor="#78909C";
            fontsize=12;
            simulator [label="IoT 数据模拟器\\n(CSV 传感器数据)", fillcolor="#ECEFF1"];
        }

        simulator -> flask [label="上传数据"];
        frontend -> flask [label="HTTP API 请求"];
        flask -> web3py [label="调用"];
        flask -> ipfs_client [label="调用"];
        web3py -> contract [label="交易/查询"];
        contract -> hardhat [label="部署在"];
        ipfs_client -> ipfs [label="存取文件"];
    }
    """
    dot_path = os.path.join(OUTPUT_DIR, "architecture.dot")
    img_path = os.path.join(OUTPUT_DIR, "architecture.png")
    with open(dot_path, "w", encoding="utf-8") as f:
        f.write(dot_src)
    subprocess.run(["dot", "-Tpng", "-Gdpi=200", dot_path, "-o", img_path], check=True)
    os.remove(dot_path)
    return img_path


def generate_workflow_diagram():
    """生成核心工作流程图"""
    dot_src = """
    digraph workflow {
        rankdir=LR;
        fontname="Microsoft YaHei";
        node [fontname="Microsoft YaHei", shape=box, style="rounded,filled", fontsize=10, width=1.6];
        edge [fontname="Microsoft YaHei", fontsize=9];
        bgcolor="white";

        start [label="IoT 设备\\n采集数据", fillcolor="#ECEFF1", shape=ellipse];
        upload [label="上传数据\\n到 IPFS", fillcolor="#E3F2FD"];
        hash [label="获取\\nIPFS Hash", fillcolor="#E3F2FD"];
        chain [label="链上存证\\n(智能合约)", fillcolor="#FFF3E0"];
        view [label="前端查看\\n数据记录", fillcolor="#E8F5E9"];
        verify [label="完整性校验\\n(哈希对比)", fillcolor="#FCE4EC"];

        start -> upload [label="CSV 文件"];
        upload -> hash [label="内容哈希"];
        hash -> chain [label="Hash+设备ID"];
        chain -> view [label="区块确认"];
        view -> verify [label="点击校验"];
    }
    """
    dot_path = os.path.join(OUTPUT_DIR, "workflow.dot")
    img_path = os.path.join(OUTPUT_DIR, "workflow.png")
    with open(dot_path, "w", encoding="utf-8") as f:
        f.write(dot_src)
    subprocess.run(["dot", "-Tpng", "-Gdpi=200", dot_path, "-o", img_path], check=True)
    os.remove(dot_path)
    return img_path


def generate_deploy_diagram():
    """生成部署架构图"""
    dot_src = """
    digraph deploy {
        rankdir=TB;
        fontname="Microsoft YaHei";
        node [fontname="Microsoft YaHei", shape=box, style="rounded,filled", fontsize=10];
        edge [fontname="Microsoft YaHei", fontsize=9];
        bgcolor="white";

        subgraph cluster_docker {
            label="Docker 容器";
            style="dashed,rounded";
            color="#2196F3";
            fontcolor="#2196F3";
            fontsize=12;

            nginx [label="Nginx\\n端口: 5140\\n(前端静态文件 + API反向代理)", fillcolor="#E3F2FD"];
            flask [label="Flask 后端\\n端口: 5141\\n(REST API 服务)", fillcolor="#E8F5E9"];
            hardhat [label="Hardhat Node\\n端口: 5145\\n(以太坊本地链)", fillcolor="#FFF3E0"];
            storage [label="本地文件存储\\n(IPFS 模拟)", fillcolor="#F3E5F5"];
        }

        user [label="用户浏览器", fillcolor="#FFFDE7", shape=ellipse];
        user -> nginx [label="HTTP :5140"];
        nginx -> flask [label="代理 /api/"];
        flask -> hardhat [label="Web3 RPC"];
        flask -> storage [label="文件读写"];
    }
    """
    dot_path = os.path.join(OUTPUT_DIR, "deploy.dot")
    img_path = os.path.join(OUTPUT_DIR, "deploy.png")
    with open(dot_path, "w", encoding="utf-8") as f:
        f.write(dot_src)
    subprocess.run(["dot", "-Tpng", "-Gdpi=200", dot_path, "-o", img_path], check=True)
    os.remove(dot_path)
    return img_path


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return heading


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def build_document():
    arch_img = generate_architecture_diagram()
    workflow_img = generate_workflow_diagram()
    deploy_img = generate_deploy_diagram()

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IoT 区块链数据存证系统")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("使 用 手 册")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("基于以太坊智能合约 + IPFS 的物联网数据安全存证平台")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph("")
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run("版本: v1.0.0    日期: 2026年3月")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ========== 目录 ==========
    add_heading(doc, "目录", level=1)
    toc_items = [
        "1. 系统概述",
        "2. 系统架构",
        "   2.1 架构分层说明",
        "   2.2 系统架构图",
        "3. 核心工作流程",
        "   3.1 流程说明",
        "   3.2 工作流程图",
        "4. 环境要求与安装",
        "   4.1 环境要求",
        "   4.2 依赖安装",
        "   4.3 启动系统",
        "   4.4 停止系统",
        "5. Docker 部署",
        "   5.1 部署架构图",
        "   5.2 构建与启动",
        "6. 功能操作指南",
        "   6.1 设备管理",
        "   6.2 数据上传与存证",
        "   6.3 数据记录查看",
        "   6.4 数据下载",
        "   6.5 数据完整性校验",
        "   6.6 权限管理",
        "7. API 接口说明",
        "8. 智能合约说明",
        "9. 常见问题",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ========== 1. 系统概述 ==========
    add_heading(doc, "1. 系统概述", level=1)
    doc.add_paragraph(
        "IoT 区块链数据存证系统是一个基于以太坊智能合约和 IPFS 分布式存储技术的物联网数据安全存证平台。"
        "系统旨在解决物联网场景下数据真实性、完整性和不可篡改性的问题。"
    )
    doc.add_paragraph("系统核心功能包括：")
    features = [
        "IoT 设备注册与管理：支持注册多种类型的传感器设备",
        "数据上传与 IPFS 存储：通过内容寻址方式存储原始数据文件",
        "区块链链上存证：将数据哈希、设备 ID、时间戳永久记录在以太坊区块链上",
        "数据完整性校验：通过对比链上哈希与本地重算哈希，检测数据是否被篡改",
        "细粒度权限控制：管理员可授权或撤销用户对特定设备数据的访问权限",
        "数据可视化仪表盘：基于 React + Ant Design 的前端界面，直观展示设备和数据信息",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ========== 2. 系统架构 ==========
    add_heading(doc, "2. 系统架构", level=1)

    add_heading(doc, "2.1 架构分层说明", level=2)
    layers = [
        ["数据层 (IPFS)", "负责存储原始的 IoT 大数据文件。由于 IPFS 是内容寻址，文件一旦上传就会生成唯一的内容哈希 (Content Hash)"],
        ["逻辑层 (Ethereum Smart Contracts)", "负责权限管理和元数据存储。智能合约不存原始数据，只存 IPFS 的哈希值，实现数据存储与验证解耦"],
        ["控制层 (Flask Backend)", "作为中间件，协调 IoT 设备、IPFS 节点和区块链之间的通信，提供 REST API 接口"],
        ["表现层 (React Frontend)", "用户查看数据和监控设备的可视化仪表盘，基于 Ant Design 组件库构建"],
    ]
    add_table(doc, ["层次", "说明"], layers)

    doc.add_paragraph("")
    add_heading(doc, "2.2 系统架构图", level=2)
    doc.add_picture(arch_img, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 3. 核心工作流程 ==========
    add_heading(doc, "3. 核心工作流程", level=1)

    add_heading(doc, "3.1 流程说明", level=2)
    steps = [
        ("数据采集与上传", "IoT 节点（用 Python 脚本模拟）读取 CSV 传感器数据集，将数据打包上传到 IPFS 节点，获取返回的 Hash_ID。"),
        ("链上存证", "后端通过 Web3.py 调用部署在 Hardhat 本地链上的智能合约，将 Hash_ID 与设备 ID、时间戳绑定，永久记录在区块链上。"),
        ("安全访问控制", "当用户请求数据时，合约首先检查该用户是否有权限。验证通过后，前端获取 Hash_ID 并从 IPFS 下载原始数据。"),
        ("完整性校验", "系统自动计算下载文件的哈希，并与链上记录的哈希对比。如果一致，证明数据未被篡改。"),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"步骤 {i}：{title}")
        run.bold = True
        doc.add_paragraph(desc)

    add_heading(doc, "3.2 工作流程图", level=2)
    doc.add_picture(workflow_img, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 4. 环境要求与安装 ==========
    add_heading(doc, "4. 环境要求与安装", level=1)

    add_heading(doc, "4.1 环境要求", level=2)
    env_reqs = [
        ["Node.js", "v18.0+", "区块链工具链和前端构建"],
        ["Python", "v3.9+", "后端 Flask 服务"],
        ["npm/pnpm", "最新版", "包管理器"],
        ["Docker", "v20.0+（可选）", "容器化部署"],
    ]
    add_table(doc, ["工具", "版本要求", "用途"], env_reqs)

    doc.add_paragraph("")
    add_heading(doc, "4.2 依赖安装", level=2)
    doc.add_paragraph("在项目根目录执行安装脚本：")
    p = doc.add_paragraph()
    run = p.add_run("./install.sh")
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xD3, 0x26, 0x03)
    doc.add_paragraph("安装脚本会依次完成以下操作：")
    install_steps = [
        "检查 Node.js、Python、npm 等基础工具",
        "安装区块链层依赖 (Hardhat + Solidity 编译器)",
        "安装前端依赖 (React + Ant Design + Vite)",
        "创建 Python 虚拟环境并安装后端依赖 (Flask + Web3.py)",
    ]
    for s in install_steps:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph("")
    add_heading(doc, "4.3 启动系统", level=2)
    doc.add_paragraph("执行一键启动脚本：")
    p = doc.add_paragraph()
    run = p.add_run("./start.sh")
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xD3, 0x26, 0x03)
    doc.add_paragraph("启动顺序如下：")
    start_steps = [
        ["步骤 1", "启动 Hardhat 本地以太坊节点", "端口 5145"],
        ["步骤 2", "编译并部署智能合约到本地链", "-"],
        ["步骤 3", "启动 Flask 后端 API 服务", "端口 5141"],
        ["步骤 4", "启动 Vite 前端开发服务器", "端口 5140"],
    ]
    add_table(doc, ["步骤", "操作", "端口"], start_steps)

    doc.add_paragraph("")
    doc.add_paragraph("启动完成后，在浏览器访问 http://localhost:5140 即可使用系统。")

    add_heading(doc, "4.4 停止系统", level=2)
    p = doc.add_paragraph()
    run = p.add_run("./stop.sh")
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xD3, 0x26, 0x03)
    doc.add_paragraph("该脚本会按反向顺序停止所有服务：前端 -> 后端 -> 区块链节点。")

    doc.add_page_break()

    # ========== 5. Docker 部署 ==========
    add_heading(doc, "5. Docker 部署", level=1)

    add_heading(doc, "5.1 部署架构图", level=2)
    doc.add_picture(deploy_img, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    add_heading(doc, "5.2 构建与启动", level=2)
    doc.add_paragraph("使用 Docker 进行一键部署，所有服务运行在同一容器中：")
    cmds = [
        ("构建镜像", "docker build -t iot-blockchain-app ."),
        ("启动容器", "docker run -d --name iot-blockchain -p 5140:5140 iot-blockchain-app:latest"),
        ("查看日志", "docker logs -f iot-blockchain"),
        ("停止容器", "docker stop iot-blockchain"),
    ]
    for desc, cmd in cmds:
        p = doc.add_paragraph()
        run = p.add_run(f"{desc}：")
        run.bold = True
        run = p.add_run(f"  {cmd}")
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xD3, 0x26, 0x03)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Docker 容器内部启动流程：Hardhat 本地链启动 -> 智能合约编译部署 -> "
        "Flask 后端启动 -> Nginx 前端服务启动。对外仅暴露 5140 端口，"
        "通过 Nginx 反向代理统一访问前端页面和后端 API。"
    )

    doc.add_page_break()

    # ========== 6. 功能操作指南 ==========
    add_heading(doc, "6. 功能操作指南", level=1)

    add_heading(doc, "6.1 设备管理", level=2)
    doc.add_paragraph("进入系统首页（设备管理页），可以看到以下内容：")
    doc.add_paragraph("设备总数、数据记录总数、设备类型数的统计卡片", style='List Bullet')
    doc.add_paragraph("IoT 设备列表表格，显示设备 ID、名称、类型、数据记录数、所有者、注册时间", style='List Bullet')
    doc.add_paragraph("")
    doc.add_paragraph("注册新设备操作步骤：")
    reg_steps = [
        '点击页面右上角的"注册设备"按钮',
        "在弹窗表单中填写设备 ID、设备名称，选择设备类型（temperature/humidity/pressure）",
        '点击"确定"提交，系统会调用智能合约将设备信息记录到区块链上',
        "注册成功后，设备列表会自动刷新显示新设备",
    ]
    for i, s in enumerate(reg_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    add_heading(doc, "6.2 数据上传与存证", level=2)
    doc.add_paragraph("数据上传支持两种方式：")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("方式一：前端界面上传")
    run.bold = True
    ui_steps = [
        '在设备列表中点击目标设备的"查看数据"按钮',
        '在数据记录页面点击"上传数据"按钮',
        "选择要上传的 CSV 数据文件",
        "系统自动完成：文件上传到 IPFS -> 获取哈希 -> 链上存证 -> 返回区块号",
    ]
    for i, s in enumerate(ui_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("方式二：IoT 模拟器批量上传")
    run.bold = True
    doc.add_paragraph(
        "运行 simulator/iot_simulator.py 脚本，可自动读取 CSV 数据集，"
        "按设备分组批量注册设备并上传数据，模拟真实 IoT 场景。"
    )

    add_heading(doc, "6.3 数据记录查看", level=2)
    doc.add_paragraph("在设备数据记录页面可以查看该设备的所有上传记录，包括：")
    record_fields = [
        ["#", "记录序号"],
        ["IPFS Hash", "文件在 IPFS 中的内容哈希（可复制）"],
        ["上传时间", "数据上链的时间戳"],
        ["区块号", "交易所在的区块编号"],
        ["上传者", "执行上传操作的以太坊账户地址"],
    ]
    add_table(doc, ["字段", "说明"], record_fields)
    doc.add_paragraph("")
    doc.add_paragraph('点击"详情"按钮可查看完整的链上存证信息和文件元数据。')

    add_heading(doc, "6.4 数据下载", level=2)
    doc.add_paragraph(
        '在数据记录列表或详情弹窗中，点击"下载"按钮即可通过 IPFS Hash 从存储系统下载原始数据文件。'
        "下载的文件内容与上传时完全一致。"
    )

    add_heading(doc, "6.5 数据完整性校验", level=2)
    doc.add_paragraph("这是系统的核心安全功能。操作步骤：")
    verify_steps = [
        '在数据记录列表中，找到要校验的记录，点击"校验"按钮',
        "系统自动执行以下操作：",
        "  a. 从 IPFS/本地存储下载文件",
        "  b. 计算文件内容的 SHA-256 哈希",
        "  c. 生成本地内容哈希并与链上记录的哈希对比",
        "校验结果弹窗显示：",
    ]
    for s in verify_steps:
        doc.add_paragraph(s)

    verify_results = [
        ["校验通过", "绿色对勾图标", "链上哈希与本地哈希完全一致，数据未被篡改"],
        ["校验失败", "红色叉号图标", "哈希不一致，数据可能已被篡改，显示红色警告"],
    ]
    add_table(doc, ["结果", "显示", "含义"], verify_results)
    doc.add_paragraph("")
    doc.add_paragraph("校验详情包括：链上哈希、本地哈希、文件 SHA-256、文件大小、链上区块号、存证时间。")

    add_heading(doc, "6.6 权限管理", level=2)
    doc.add_paragraph("管理员可以对设备数据进行细粒度的访问控制：")
    perm_ops = [
        '在设备列表中点击"权限管理"按钮进入权限管理页面',
        "查看当前设备的已授权用户列表",
        '授权新用户：输入以太坊地址，点击"授权"按钮',
        '撤销权限：在已授权列表中点击"撤销"按钮',
        "所有权限变更都通过智能合约执行，记录在区块链上不可篡改",
    ]
    for i, s in enumerate(perm_ops, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_page_break()

    # ========== 7. API 接口说明 ==========
    add_heading(doc, "7. API 接口说明", level=1)
    doc.add_paragraph("后端提供 REST API 接口，基础地址：http://localhost:5140/api")
    doc.add_paragraph("")

    apis = [
        ["GET", "/api/health", "健康检查", "返回服务状态和区块链连接信息"],
        ["GET", "/api/devices", "获取设备列表", "返回所有已注册设备及其信息"],
        ["POST", "/api/devices", "注册新设备", "请求体: {deviceId, deviceName, deviceType}"],
        ["POST", "/api/upload", "上传数据文件", "FormData: file + deviceId，返回 IPFS Hash"],
        ["GET", "/api/records?deviceId=xxx", "查询设备记录", "返回指定设备的所有链上记录"],
        ["GET", "/api/download?hash=xxx", "下载数据文件", "通过 IPFS Hash 下载原始文件"],
        ["POST", "/api/verify", "完整性校验", "请求体: {deviceId, ipfsHash}，返回校验结果"],
        ["POST", "/api/access/grant", "授权用户", "请求体: {deviceId, userAddress}"],
        ["POST", "/api/access/revoke", "撤销授权", "请求体: {deviceId, userAddress}"],
        ["GET", "/api/access/check", "检查权限", "查询参数: deviceId, userAddress"],
    ]
    add_table(doc, ["方法", "路径", "功能", "说明"], apis)

    doc.add_page_break()

    # ========== 8. 智能合约说明 ==========
    add_heading(doc, "8. 智能合约说明", level=1)
    doc.add_paragraph(
        "系统使用 Solidity 编写的 IoTDataStorage 智能合约，部署在以太坊网络上（开发环境使用 Hardhat 本地链）。"
    )
    doc.add_paragraph("")

    add_heading(doc, "核心功能函数", level=2)
    contract_funcs = [
        ["registerDevice", "注册 IoT 设备", "仅管理员可调用"],
        ["storeData", "存储数据记录", "记录 IPFS Hash、设备 ID 和时间戳"],
        ["getDataByDevice", "按设备查询数据", "需要访问权限"],
        ["grantAccess", "授权用户访问", "仅管理员可调用"],
        ["revokeAccess", "撤销用户权限", "仅管理员可调用"],
        ["checkAccess", "检查访问权限", "返回布尔值"],
        ["verifyData", "验证数据存在性", "检查哈希是否在链上"],
    ]
    add_table(doc, ["函数名", "功能", "权限要求"], contract_funcs)

    doc.add_paragraph("")
    doc.add_paragraph("合约特性：")
    contract_features = [
        "采用 OpenZeppelin 风格的权限控制模式",
        "管理员为合约部署者（Hardhat 第一个账户）",
        "所有数据存储操作产生事件日志，便于链下监听",
        "支持按设备 ID 索引的数据查询",
    ]
    for f in contract_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ========== 9. 常见问题 ==========
    add_heading(doc, "9. 常见问题", level=1)

    faqs = [
        (
            "Q: 启动时提示端口被占用怎么办？",
            "A: 先执行 ./stop.sh 停止旧服务，或手动释放对应端口（5140/5141/5145）。"
        ),
        (
            "Q: 前端页面空白或显示错误？",
            "A: 确保使用 NODE_ENV=development 启动 Vite，或重启前端服务：kill 掉旧进程后重新运行 npx vite --port 5140。"
        ),
        (
            "Q: 区块链连接失败？",
            "A: 检查 Hardhat 节点是否正常运行在 5145 端口。可查看 logs/hardhat.log 获取详细信息。"
        ),
        (
            "Q: 数据校验失败但文件未被修改？",
            "A: 确认 IPFS/本地存储的文件未被意外修改。校验机制通过重新计算文件 SHA-256 与链上哈希对比，任何字节级改动都会导致校验失败。"
        ),
        (
            "Q: Docker 部署后如何查看内部日志？",
            "A: 执行 docker logs -f <容器名> 查看 Nginx 日志，或 docker exec -it <容器名> cat /var/log/supervisor/flask.log 查看后端日志。"
        ),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        doc.add_paragraph(a)
        doc.add_paragraph("")

    # 保存文档
    output_path = os.path.join(OUTPUT_DIR, "IoT区块链数据存证系统-使用手册.docx")
    doc.save(output_path)
    print(f"使用手册已生成: {output_path}")

    # 清理临时图片
    for img in [arch_img, workflow_img, deploy_img]:
        if os.path.exists(img):
            os.remove(img)

    return output_path


if __name__ == "__main__":
    build_document()
