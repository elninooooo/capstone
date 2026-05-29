"""Insert Section 4.2 Testing into Capstone_Example.docx.

What this does:
  1. Finds the paragraph styled like a section header named "Results and Discussion".
  2. Inserts BEFORE it: "Testing" section heading + 7 subsections of prose,
     with figures (numbered Figure 6 through Figure 13) and a single new table
     (Table 9: Test Coverage Summary).
  3. Adds the new figures + table to the manual List of Figures / List of Tables.

Heading conventions reused from the existing document:
  - Section heading           -> 'Body Text' style (matches "Implementation"
                                  and "Results and Discussion" in the original)
  - Subsection heading        -> 'Heading 3' style
  - List-of-figures entries   -> 'Heading 5' style
  - Body prose, captions      -> 'Body Text' style

The script is idempotent-safe: it refuses to run if the new "Testing" heading
is already present.

Output: Capstone_Example.docx is rewritten in place. A backup was made
manually beforehand (Capstone_Example.before-testing-insert.docx).
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "Capstone_Example.docx"
FIG_DIR = ROOT / "figures_testing"

# --- New figure / table numbering (continuing the existing series) -----------
FIG_BASE = 5      # last existing Figure N is 5
TBL_BASE = 8      # last existing Table N in body is 8

FIGURES = [
    # (T-id,  filename,                                      caption text)
    ("T1", "T1_test_pyramid.png",
        "Test Classification Pyramid (bar widths proportional to test count)"),
    ("T2", "T2_contract_distribution.png",
        "Smart Contract Test-Case Distribution by Module (27 cases / 7 suites)"),
    ("T3", "test-contract-access-granted.png",
        "Smart contract access-grant transaction (Hardhat console output)"),
    ("T4", "test-api-upload-result.png",
        "Successful upload response with IPFS CID and blockchain tx hash"),
    ("T5", "test-api-unauthorized.png",
        "Backend RBAC denial — HTTP 403 Forbidden response"),
    ("T6", "test-iot-simulator-summary.png",
        "IoT simulator end-to-end run summary"),
    ("T7", "test-frontend-access-granted.png",
        "Dashboard view confirming an access grant"),
    ("T8", "test-verify-fail-tampered.png",
        "Tamper-detection verification failure (red \"Integrity Compromised\" modal)"),
]

# Map T-id -> rendered "Figure N" once we know the base.
def fig_no(tid: str) -> int:
    return FIG_BASE + 1 + [f[0] for f in FIGURES].index(tid)


# =============================================================================
# Low-level helpers
# =============================================================================

def insert_paragraph_before(ref_para, text: str, *, style: str = None,
                            bold: bool = False, italic: bool = False) -> "Paragraph":
    """Create a new paragraph immediately before `ref_para` and return it."""
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, ref_para._parent)
    if style:
        para.style = ref_para._parent.part.document.styles[style]
    if text:
        run = para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return para


def insert_picture_before(ref_para, image_path: Path, *, width_inches: float = 5.5) -> None:
    """Add a centred picture in a new paragraph just before `ref_para`."""
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    para = Paragraph(new_p, ref_para._parent)
    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def insert_table_before(ref_para, headers, rows, *, style: str = "Table Grid"):
    """Insert a docx table with `headers` row + `rows` of cells before ref_para."""
    doc = ref_para._parent.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = doc.styles[style]
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    # Move the table element to immediately before ref_para
    ref_para._p.addprevious(table._tbl)
    return table


# =============================================================================
# Content writers (each takes ref_para = the "Results and Discussion" para)
# =============================================================================

def write_section_intro(ref):
    insert_paragraph_before(ref, "Testing", style="Body Text", bold=True)
    insert_paragraph_before(
        ref,
        "Before presenting the empirical results in the next section, this section "
        "describes the testing methodology, the test environment, and the test "
        "artefacts produced during validation of the prototype. The purpose of "
        "separating testing from results is to make explicit what was tested, "
        "how it was tested, and which evidence was collected, so that the discussion "
        "that follows can focus on the interpretation of those results rather than "
        "on the test design itself.",
        style="Body Text",
    )


def write_subsection(ref, title: str, paragraphs: list[str]):
    insert_paragraph_before(ref, title, style="Heading 3")
    for p in paragraphs:
        insert_paragraph_before(ref, p, style="Body Text")


def write_figure(ref, tid: str):
    fname, cap = next((f, c) for tt, f, c in FIGURES if tt == tid)
    insert_picture_before(ref, FIG_DIR / fname, width_inches=5.5)
    n = fig_no(tid)
    cap_para = insert_paragraph_before(ref, f"Figure {n}: {cap}", style="Body Text", italic=True)
    cap_para.alignment = 1  # centred caption


# =============================================================================
# Main insert
# =============================================================================

def main():
    doc = Document(DOCX)

    # --- locate the section header "Results and Discussion" (Body Text, exact match)
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Results and Discussion" and p.style.name == "Body Text":
            target = p
            break
    if target is None:
        raise SystemExit("Could not find 'Results and Discussion' Body Text paragraph.")

    # idempotency guard
    for p in doc.paragraphs:
        if p.text.strip() == "Testing" and p.style.name == "Body Text":
            raise SystemExit("Document already contains a 'Testing' section header.")

    # ---- Section intro ----
    write_section_intro(target)

    # ---- 4.2.1 Test Strategy and Classification ----
    write_subsection(target, "Test Strategy and Classification", [
        "Validation followed a three-tier strategy that mirrors the layered architecture of "
        "the framework. At the lowest layer, the on-chain logic was covered by a comprehensive "
        "Solidity unit-test suite executed in the Hardhat environment, which is the de-facto "
        "practice for Ethereum-based contract development [11]. At the middle layer, the Flask "
        "backend was exercised through HTTP-level integration tests issued from curl and Postman "
        "against a running development server. At the top layer, the React dashboard and the "
        "Python IoT simulator were used to drive exploratory end-to-end scenarios that traverse "
        "all four system layers (smart contract → backend → off-chain storage → frontend). For "
        "the prototype scope, automated test suites were written only for the smart-contract "
        "layer; the backend and frontend layers were validated through reproducible manual "
        "scenarios with screenshot evidence, which is consistent with the rapid-prototyping "
        "approach reported in comparable blockchain–IoT studies [1] [13]. The distribution of "
        f"test cases across these three tiers is summarised in Figure {fig_no('T1')}.",
    ])
    write_figure(target, "T1")

    # ---- 4.2.2 Test Environment and Dataset ----
    write_subsection(target, "Test Environment and Dataset", [
        "All tests were executed on the development workstation and the four-port runtime "
        "described in Section 4.1.1 (frontend on 5140, backend on 5141, blockchain on 5145, "
        "IPFS on 5146). The smart-contract test runner uses Hardhat's in-memory Ethereum "
        "network, which provides a fresh chain state for every test through the beforeEach "
        "hook, eliminating cross-test contamination. The integration and end-to-end tests "
        "were executed against a live system started with the start.sh script, which "
        "sequentially launches the Hardhat node, deploys the contract, starts the Flask "
        "backend, and serves the Vite frontend.",

        "A single canonical dataset (simulator/data/iot_sensor_data.csv) was used as the "
        "primary fixture across all test tiers. It contains 24 sensor records distributed "
        "evenly across three devices—TEMP-001 (8 records, temperature), HUM-002 (8 records, "
        "humidity), and PRESS-003 (8 records, pressure). With the default batch size of "
        "four records per upload, this fixture produces six on-chain attestation transactions, "
        "allowing every workflow assertion in the next section to be reproduced "
        "deterministically. Two auxiliary fixtures (test_sensor_data.csv and "
        "test_sensor_data2.csv) were prepared specifically for the tamper-detection "
        "experiment described in §4.2.6: the former is the original file recorded on chain, "
        "the latter is a structurally compatible but content-divergent file used to simulate "
        "an off-chain modification.",
    ])

    # ---- 4.2.3 Smart Contract Unit Testing ----
    write_subsection(target, "Smart Contract Unit Testing", [
        "The Solidity contract IoTDataStorage.sol is covered by the test file "
        "blockchain/test/IoTDataStorage.test.js, executed via npx hardhat test. The suite "
        "contains 27 individual test cases organised into 7 functional describe blocks that "
        "map one-to-one to the contract modules introduced in §4.1.2: Deployment, Device "
        "Registration, Data Storage, Data Query, Access Control, Data Verification, and "
        "Admin Transfer. The per-module distribution of test cases is shown in "
        f"Figure {fig_no('T2')}.",
    ])
    write_figure(target, "T2")

    write_subsection(target, "", [
        "Each test case belongs to one of four assertion categories that together provide "
        "functional, security, and structural coverage of the contract: state assertions "
        "verify that storage mappings (devices, records, accessPermissions) and counters "
        "reflect the expected values after a transaction; event assertions use Chai's "
        "to.emit(...).withArgs(...) matcher to confirm that DeviceRegistered, DataStored, "
        "AccessGranted, and AccessRevoked are emitted with the correct payloads, which is "
        "essential because the frontend relies on these events for near-real-time UI updates; "
        "revert assertions validate the negative paths enforced by the onlyAdmin, deviceExists, "
        "and hasAccess modifiers (e.g. \"Only admin can perform this action\", \"Device not "
        "registered\", \"No access permission for this device\"); and boundary assertions "
        "check edge inputs such as out-of-bounds record indices and the zero-address "
        "rejection in transferAdmin.",

        "The Hardhat optimiser (200 runs) is enabled during compilation, so the deployed "
        "bytecode used during testing matches the bytecode that would be deployed in "
        f"production. Figure {fig_no('T3')} shows a representative on-chain interaction "
        "recorded during these tests—the AccessGranted event being emitted after a "
        "successful grantAccess call.",
    ])
    write_figure(target, "T3")

    # ---- 4.2.4 Backend API Integration Testing ----
    write_subsection(target, "Backend API Integration Testing", [
        "The Flask backend exposes 13 REST endpoints (§4.1.3) organised into five functional "
        "groups: device management (/api/devices, /api/devices/<id>), data movement "
        "(/api/upload, /api/download), data query (/api/records, /api/records/<id>/<index>), "
        "integrity verification (/api/verify), and access control (/api/access/grant, "
        "/api/access/revoke, /api/access/check). Two diagnostic endpoints (/api/health, "
        "/api/blockchain/accounts) round out the surface.",

        "Because the backend's primary responsibility is orchestration—translating HTTP "
        "requests into combined IPFS-client and blockchain-client operations—it was validated "
        "through reproducible HTTP-level scenarios issued from the command line. Each "
        "endpoint was exercised through both its happy path and its principal failure paths, "
        "with the relevant HTTP status codes asserted by inspection (201 on create, 400 on "
        "invalid input, 403 on access denial, 404 on missing resource, 503 when the "
        f"blockchain client is unreachable). Figure {fig_no('T4')} shows the JSON response "
        "of a successful upload, including the IPFS CID, the SHA-256 checksum, and the "
        f"simulated blockchain transaction hash; Figure {fig_no('T5')} shows the corresponding "
        "403 response for an unauthorised access attempt.",
    ])
    write_figure(target, "T4")
    write_figure(target, "T5")

    # ---- 4.2.5 Frontend and End-to-End Testing ----
    write_subsection(target, "Frontend and End-to-End Testing", [
        "The React dashboard was validated through scripted manual walkthroughs that "
        "exercise each of the five primary routes defined in §4.1.4: device list, "
        "per-device record browsing, the integrity-verification modal embedded in the "
        "record view, and the access-management page. Each walkthrough was recorded with "
        "screenshots of the relevant UI states (registration confirmation, upload progress, "
        "verification verdict, grant / revoke confirmations).",

        "For full end-to-end validation, the IoT simulator (simulator/iot_simulator.py) was "
        "used to drive the entire system from data ingestion to on-chain attestation in a "
        "single automated run. The simulator was invoked with default arguments "
        "(--csv data/iot_sensor_data.csv --batch-size 4 --api-url http://127.0.0.1:5141) "
        "and executed its three sequential phases: (i) loading and grouping the 24-record "
        "CSV by device, (ii) registering each unique device through POST /api/devices, and "
        "(iii) batching every device's records into chunks of four and uploading each chunk "
        "through POST /api/upload. The simulator's per-batch log lines—reporting the IPFS "
        "CID, block number, and batch size—and its terminal summary serve as primary evidence "
        f"for the workflow assertions in the next section. Figure {fig_no('T6')} reproduces "
        f"this terminal summary, and Figure {fig_no('T7')} shows the corresponding state of "
        "the dashboard's access-management view immediately after a permission grant.",
    ])
    write_figure(target, "T6")
    write_figure(target, "T7")

    # ---- 4.2.6 Tamper-Detection Testing ----
    write_subsection(target, "Tamper-Detection Testing", [
        "The tamper-detection capability of the hash-based integrity mechanism was validated "
        "through a controlled experiment. After a record had been uploaded and attested on "
        "chain, the corresponding off-chain file in backend/uploads/ was replaced with a "
        "content-divergent file (test_sensor_data2.csv) of the same name. A subsequent "
        "POST /api/verify request for the original IPFS hash and device identifier was "
        "issued, and the response was inspected. The expected behaviour is that the "
        "chain-recorded hash and the recomputed SHA-256 of the on-disk file disagree, "
        "causing the backend to return verified: false and the frontend to display a red "
        f"\"Data Integrity Compromised\" indicator (Figure {fig_no('T8')}). After the "
        "experiment, the original file was restored and a follow-up verification was issued; "
        "the system returned to the \"Integrity Verified\" state, confirming that the "
        "verification logic is purely a function of the file content and not of any cached "
        "state.",
    ])
    write_figure(target, "T8")

    # ---- 4.2.7 Test Coverage Summary ----
    write_subsection(target, "Test Coverage Summary", [
        f"Table {TBL_BASE + 1} summarises the per-layer test coverage achieved by the "
        "strategy described above, the type of evidence collected, and the principal "
        "limitations carried into the discussion that follows.",
    ])

    insert_paragraph_before(
        target,
        f"Table {TBL_BASE + 1}: Test Coverage Summary",
        style="Body Text", bold=True,
    )

    insert_table_before(target,
        headers=["Layer", "Test type", "Test count", "Automation",
                 "Primary evidence", "Known limitation"],
        rows=[
            ["Smart contract", "Unit", "27 cases / 7 suites", "Automated (Hardhat)",
             "Mocha pass output", "No Gas profiling"],
            ["Backend API", "Integration (HTTP)", "13 endpoints, both paths",
             "Manual (curl / Postman)", "JSON responses + screenshots",
             "No load / concurrency testing"],
            ["Frontend / E2E", "Exploratory", "5 routes + simulator run",
             "Manual", "Walkthrough screenshots",
             "No browser automation (e.g. Playwright)"],
            ["Security", "Tamper-detection", "1 controlled experiment",
             "Manual", "Verification modal screenshots", "Single mutation type only"],
        ],
    )

    insert_paragraph_before(
        target,
        "These coverage choices are appropriate to a single-developer prototype and "
        "exercise every architectural feature claimed in Chapter 3, but they leave three "
        "gaps that are revisited in the discussion: absence of on-chain Gas measurement, "
        "absence of load testing, and absence of browser-driven UI automation. With the "
        "testing methodology established, the following section now reports the empirical "
        "results obtained from executing this test plan.",
        style="Body Text",
    )

    # =====================================================================
    # Update the manual List of Figures + List of Tables
    # Use *fresh* paragraph scans after each insertion, because earlier inserts
    # invalidate any cached indices.
    # =====================================================================
    def find_anchor_after(text_match: str, start_text: str) -> "Paragraph | None":
        """Return the paragraph immediately after the one whose text starts
        with `start_text`, scanning fresh from `text_match` ('List of …')."""
        in_section = False
        prev = None
        for p in doc.paragraphs:
            if p.text.strip() == text_match:
                in_section = True
                continue
            if in_section and prev is not None and prev.text.strip().startswith(start_text):
                return p
            prev = p
        return None

    def find_first_nonmatching_after(header_text: str, prefix: str) -> "Paragraph | None":
        """Return the first paragraph after the header whose text does NOT
        start with `prefix`."""
        in_section = False
        for p in doc.paragraphs:
            if p.text.strip() == header_text:
                in_section = True
                continue
            if in_section and not p.text.strip().startswith(prefix):
                # skip blank paragraphs immediately after header
                if p.text.strip() or in_section_started:  # type: ignore
                    pass
                return p
        return None

    # Insert new Figure entries in List of Figures, after the last "Figure N" entry
    lof_anchor = find_anchor_after("List of Figures", "Figure 5")
    if lof_anchor is not None:
        for tid, _, cap in FIGURES:
            entry = f"Figure {fig_no(tid)}: {cap}\t—"
            insert_paragraph_before(lof_anchor, entry, style="Heading 5")

    # Insert new Table entry in List of Tables, after the last "Table N" entry.
    # Find the LAST paragraph that starts with "Table " in the List of Tables,
    # then insert before the paragraph after it.
    in_lot = False
    last_table_para = None
    after_last_table = None
    for p in doc.paragraphs:
        if p.text.strip() == "List of Tables":
            in_lot = True
            continue
        if not in_lot:
            continue
        if p.text.strip().startswith("Table "):
            last_table_para = p
        elif last_table_para is not None and p.text.strip():
            # first non-empty, non-Table paragraph after the table list
            after_last_table = p
            break
    # Use the paragraph right after the last "Table N" entry as the anchor.
    if last_table_para is not None and after_last_table is not None:
        insert_paragraph_before(
            after_last_table,
            f"Table {TBL_BASE + 1}: Test Coverage Summary\t—",
            style="Heading 5",
        )

    # =====================================================================
    doc.save(DOCX)
    print(f"OK: wrote updated {DOCX.name}")


if __name__ == "__main__":
    main()
